"""Word document operations module.

This module provides functions to read and edit Word documents using
a modular approach with separate operation handlers.
"""

import json
from .executor import execute_python_code
from .word_operations import (
    replace_text_operation,
    replace_paragraph_operation,
    insert_text_operation,
    insert_heading_operation,
    delete_paragraph_operation,
)


def read_word_content(session_id: str) -> str:
    """Đọc nội dung Word trả về text kèm index để dễ dàng chỉnh sửa.
    
    Args:
        session_id: ID của session để thực thi code
        
    Returns:
        Kết quả đọc nội dung Word
    """
    code = '''
from docx import Document
import os

# Ưu tiên lấy file _edited.docx nếu có để đọc nội dung mới nhất
docx_files = [f for f in os.listdir('/app/data') if f.endswith('.docx')]
edited_files = [f for f in docx_files if f.endswith('_edited.docx')]
original_files = [f for f in docx_files if not f.endswith('_edited.docx')]

if edited_files:
    filename = edited_files[0]
elif original_files:
    filename = original_files[0]
else:
    print("Không tìm thấy file Word!")
    exit(1)

doc = Document(f'/app/data/{filename}')
print(f"Nội dung file: {filename}\\n")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip(): 
        print(f"[{i}] {p.text}")
'''
    return execute_python_code(code, session_id)


def edit_word_document(session_id: str, operations: list) -> str:
    """Chỉnh sửa file Word (.docx) đã được tải lên bằng các lệnh cấu trúc.
    
    Args:
        session_id: ID của session để thực thi code
        operations: Danh sách các operations cần thực hiện
        
    Returns:
        Kết quả thực thi các operations
    """
    # Map operation types to their handler functions
    operation_handlers = {
        'replace': replace_text_operation,
        'replace_paragraph': replace_paragraph_operation,
        'insert_text': insert_text_operation,
        'insert_heading': insert_heading_operation,
        'delete_paragraph': delete_paragraph_operation,
    }
    
    # Build the operation code
    operations_code = []
    for op in operations:
        op_type = op.get('type')
        
        if op_type == 'custom_code':
            # Handle custom code directly
            custom_code = op.get('code')
            operations_code.append(f'''
# Custom code operation
custom_code = {repr(custom_code)}
try:
    exec(custom_code, globals())
    print(f"- Đã thực thi custom code thành công.")
except Exception as e:
    print(f"- Lỗi thực thi custom code: {{e}}")
''')
        elif op_type in operation_handlers:
            # Use the appropriate handler
            operations_code.append(operation_handlers[op_type](op))
        else:
            operations_code.append(f'''
print(f"- Cảnh báo: Operation type '{{op_type}}' không được hỗ trợ")
''')
    
    # Combine all operations into final code
    code = f'''
import json
from docx import Document
import os

# Tìm file Word: Ưu tiên file đã sửa trước đó để edit lặp lại
docx_files = [f for f in os.listdir('/app/data') if f.endswith('.docx')]
edited_files = [f for f in docx_files if f.endswith('_edited.docx')]
original_files = [f for f in docx_files if not f.endswith('_edited.docx')]

if edited_files:
    # Nếu đang sửa tiếp, lấy file đã sửa gần nhất làm gốc
    filename = edited_files[0]
    base_name = filename.replace('_edited.docx', '')
elif original_files:
    # Nếu lần đầu sửa, lấy file gốc
    filename = original_files[0]
    base_name = filename.replace('.docx', '')
else:
    print("Lỗi: Không tìm thấy file Word nào!")
    exit(1)

doc = Document(f'/app/data/{{filename}}')
print(f"Đang xử lý file: {{filename}}")

{chr(10).join(operations_code)}

# Lưu file kết quả
output_filename = f"{{base_name}}_edited.docx"
doc.save(f'/app/data/{{output_filename}}')
print(f"\\nĐã lưu file chỉnh sửa: {{output_filename}}")
'''
    return execute_python_code(code, session_id)
